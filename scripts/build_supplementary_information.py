# -*- coding: utf-8 -*-
"""Build the current Supplementary Information DOCX from validated figure bundles."""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SUP_DIR = 'results/paper_figure_multi_seed/supplementary_v5_c5_revised_20260804_r2/figures'
S7_DIR = 'results/paper_figure_multi_seed/supplementary_v5_s7_complete_pairs_20260812_r1/figures'
OUT = 'docs/paper/supplementary_information.docx'

doc = Document()
# base style
st = doc.styles['Normal']
st.font.name = 'Times New Roman'
st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13)
    return p

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11.5)
    return p

def para(text, size=10.5, bold_prefix=None, italic=False):
    p = doc.add_paragraph()
    if bold_prefix:
        rb = p.add_run(bold_prefix); rb.bold = True; rb.font.size = Pt(size)
    r = p.add_run(text); r.font.size = Pt(size); r.italic = italic
    return p

def legend_title(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(10.5)
    return p

def legend(text):
    return para(text, size=10.5)

def figure(png, width_in=6.3):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(png, width=Inches(width_in))
    return p

def table(header, rows, widths=None, font=8.5):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = 'Table Grid'
    for j, htxt in enumerate(header):
        c = t.rows[0].cells[j]
        c.text = ''
        r = c.paragraphs[0].add_run(htxt); r.bold = True; r.font.size = Pt(font)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ''
            r = cells[j].paragraphs[0].add_run(str(val)); r.font.size = Pt(font)
    return t

# ---------------- Title ----------------
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('Supplementary Information'); r.bold = True; r.font.size = Pt(16)
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run('History-Conditioned STSP Links Working-Memory Maintenance and Organization')
r.bold = True; r.font.size = Pt(13)
para('Yunpeng Xiang, Yuxi Liu and Qing Pan', size=11)
para('', size=6)

para('This file contains Supplementary Figures S1\u2013S7, Supplementary Table S1 (model, encoding and training parameters) and Supplementary Table S2 (endpoint definitions and statistical analysis plan) for the main manuscript. All analyses used the same 20 independently trained networks (seeds 1000\u20131019) as the main text; the independently trained network was the inferential unit throughout, and trial-, unit-, site- and stage-level values were aggregated within each network before any cross-network comparison. Error bars and intervals show two-sided 95% confidence intervals (Student t for Figs. S1\u2013S3, S5\u2013S7 descriptive summaries; bootstrap for the exact sign-flip endpoints in Fig. S4 and for the confirmatory transfer endpoints in Fig. S2), unless stated otherwise. Machine-readable network-level statistics and all per-network values are provided with the Source Data accompanying this manuscript.', size=10.5)

# ---------------- S1 ----------------
h2('Supplementary Figure S1 | Delay-dependent and donor-directed controls for the functional inherited-state effect')
figure(f'{SUP_DIR}/supp_fig_s1.png')
legend_title('Supplementary Fig. S1 | Delay dependence and donor-directed controls for the functional inherited-state effect in Fig. 2. ')
legend('a, Static-frozen minus dynamic-intact probe accuracy across delays from 100 to 1,200 ms. The gap declined from 23.35 percentage points (95% CI, 21.03\u201325.67) at 100 ms to 4.35 (95% CI, 2.56\u20136.14) at 1,200 ms. b, Dynamic-intact minus static-frozen sample-label prediction rate across the same delays, declining from 7.00 percentage points (95% CI, 5.88\u20138.12) to 1.15 (95% CI, 0.39\u20131.91). c, Trial-paired donor flux: inflow, outflow and net trials, showing that the aggregate donor shift in Fig. 2d reflects paired trial-level substitution (inflow 6.05% [95% CI, 5.17\u20136.93]; outflow 1.80% [1.28\u20132.32]; net 4.25% [3.23\u20135.27]). d, Donor excess calibration: excess under intact readout was indistinguishable from zero (\u22120.28 percentage points [95% CI, \u22120.88 to 0.32]), whereas trial-shuffled readout produced a positive excess (3.50 percentage points [95% CI, 2.71\u20134.29]; shuffled-minus-intact difference, 3.78 [95% CI, 2.87\u20134.69]), excluding unequal donor-label opportunity as an explanation. Means across n = 20 independently trained networks with two-sided 95% Student t CIs; all panels provide descriptive estimates.')

# ---------------- S2 ----------------
h2('Supplementary Figure S2 | Causal identity of the Layer 1-only transfer and its robustness')
figure(f'{SUP_DIR}/supp_fig_s2.png')
legend_title('Supplementary Fig. S2 | The Layer 1-only donor transfer in Fig. 4f is causally isolated, holds on the same donor-transfer axis in Layer 2 and early Layer 3, and is unchanged when the development seed is excluded. ')
legend('a, Causal-identity gate of the selective transfer: the same encoded B was presented after Layer 1 u/x exchange only, with fast state equalized, Layer 2 and Layer 3 STSP states retained and receiver validity confirmed for both readout layers. b, Donor-transfer index for the Layer 2 update and the early Layer 3 successor under this isolation: 0.8086 (95% bootstrap CI, 0.8076\u20130.8097) and 0.5210 (95% CI, 0.4661\u20130.5837); both confirmatory contrasts exceeded zero in every network (Holm-adjusted P = 7.63 \u00d7 10\u207b\u2076, confirmatory core family). c, Robustness margins above null for the two endpoints: Layer 2 alignment 0.8565 (95% CI, 0.8557\u20130.8573), median donor-transfer index 0.8145 (95% CI, 0.8131\u20130.8159) and positive fraction 0.5 (ceiling); early Layer 3 alignment 0.4647 (95% CI, 0.4507\u20130.4786), median donor-transfer index 0.5384 (95% CI, 0.5232\u20130.5537) and positive fraction 0.2329 (95% CI, 0.2253\u20130.2405). d, Untouched-cohort sensitivity excluding the development seed: Layer 2 update donor-transfer index 0.8087 (95% CI, 0.8075\u20130.8098) and early Layer 3 0.5190 (95% CI, 0.4620\u20130.5843) for the remaining n = 19 networks (Holm-adjusted P = 1.53 \u00d7 10\u207b\u2075), indistinguishable from the full cohort. Points and bars summarize n = 20 networks (n = 19 in d); intervals are two-sided 95% bootstrap CIs for transfer values and Student t CIs for margins.')

# ---------------- S3 ----------------
h2('Supplementary Figure S3 | Local transition composition is robust to window, winner cap, distance threshold and original-winner fate')
figure(f'{SUP_DIR}/supp_fig_s3.png')
legend_title('Supplementary Fig. S3 | The early-window transition and overlap-targeted perturbation effects in Fig. 4c,d are robust to analysis choices. ')
legend('a, Transition gain as a function of the early-window size (5\u201330 ms) for probe-only and random-matched comparators, showing that the overlap-related transition gain is not an artefact of a single window choice. b, Winner-minus-loser membrane-voltage difference for the top one, two and three ranked candidates (Top 1, 0.181 [95% CI, 0.099\u20130.264]; Top 2, 0.170 [0.103\u20130.237]; Top 3, 0.155 [0.099\u20130.210] mV), confirming that the effect does not depend on a single winner definition. c, The same voltage difference under distance caps of \u22642, \u22644 and \u22646 units (0.269 [95% CI, 0.226\u20130.311], 0.212 [0.168\u20130.257] and 0.177 [0.128\u20130.226] mV), ruling out a spatial-threshold artefact. d, Fate of the original winner after targeted overlap-aligned attenuation or reset: attenuation preserved the winner in 33.02% (95% CI, 32.85\u201333.19) of events, delayed it in 58.78% (58.62\u201358.94) and lost it in 8.20% (8.06\u20138.34); reset delayed the winner in 85.16% (84.94\u201385.38) and lost it in 14.84% (14.62\u201315.06), with no preserved events. The fate decomposition reports how original winners were redistributed among preserved, delayed and lost outcomes after each intervention. n = 20 networks; two-sided 95% Student t CIs; descriptive endpoints.')

# ---------------- S4 ----------------
h2('Supplementary Figure S4 | Successor-transplant donor transfer is stable across history depths and cohorts, with exact identity gates')
figure(f'{SUP_DIR}/supp_fig_s4.png')
legend_title('Supplementary Fig. S4 | The Fig. 5a,b donor-transfer results are reproducible across history depths, unchanged after excluding the development seed and gated on exact donor-receiver identity. ')
legend('a, Network-level donor-transfer index in the early Layer 2 event map at history depths K = 1 and K = 5 (0.337 [95% bootstrap CI, 0.336\u20130.339] and 0.344 [0.342\u20130.345]). b, Donor-transfer index in the post-C Layer 3 successor at K = 1 and K = 5 (0.283 [0.281\u20130.284] and 0.291 [0.290\u20130.292]). All four endpoints were positive in every network (one-sided exact sign-flip tests, Holm-adjusted P = 3.81 \u00d7 10\u207b\u2076). c, Untouched-cohort sensitivity (n = 19, development seed excluded) for the same four endpoints (K = 1 early response, 0.3373 [95% CI, 0.3361\u20130.3385]; K = 5, 0.3441 [0.3425\u20130.3456]; K = 1 successor, 0.2824 [0.2809\u20130.2839]; K = 5, 0.2906 [0.2894\u20130.2919]), matching the full cohort within confidence limits. d, Identity-gate audit: 100% of transplanted successor samples passed exact donor-receiver identity checks at every history depth (rows), confirming that the transfer statistics operate on verified donor states. n = 20 networks (n = 19 in c); intervals are two-sided 95% bootstrap CIs.')

# ---------------- S5 ----------------
h2('Supplementary Figure S5 | Observed-over-passive displacement recurs across networks and stages and is carried by both STSP variables')
figure(f'{SUP_DIR}/supp_fig_s5.png')
legend_title('Supplementary Fig. S5 | The Fig. 5c recurrence result is not driven by a few networks or stages and is expressed in both STSP variables. ')
legend('a, Network \u00d7 stage heatmap of observed-minus-passive joint-state displacement (stages 2\u201310, 20 networks), showing that input-associated displacement exceeded the equal-time passive benchmark broadly rather than in isolated cells. b, Minimum observed-minus-passive displacement across stages within each network, 0.4890 (95% CI, 0.4818\u20130.4960), positive in all 20 networks even at the worst stage. c,d, Displacement trajectories of the utilization variable u and the available-resource variable x across stages 2\u201310, showing that the recurrence is not carried by a single STSP variable. Descriptive endpoints; means across n = 20 networks with two-sided 95% Student t CIs.')

# ---------------- S6 ----------------
h2('Supplementary Figure S6 | Accumulated-state morphology shows a coefficient-free spatial boundary in Layer 1 effective support')
figure(f'{SUP_DIR}/supp_fig_s6.png')
legend_title('Supplementary Fig. S6 | The accumulated-state organization in Fig. 6c\u2013f remains structured under coefficient-free measures, ruling out NNLS coefficient artefacts and uniform fading. ')
legend('a, Effective component number (Neff) by sequence length under the raw NNLS coefficients used for display (2.99, 4.96, 6.74 and 8.05 at K = 3, 5, 7 and 10) and under a column-normalized similarity-based measure (3.00, 5.00, 7.00 and 10.00), both increasing with load. b, Latest-item weight across the same lengths (NNLS: 0.331, 0.195, 0.133 and 0.073; similarity-based: 0.335, 0.202, 0.145 and 0.101), declining with load under both measures. c, Mean no-memory-corrected effective-support change, \u0394g, across the load-by-delay grid. d, Rook-Moran excess over the same grid, declining from 0.885 (95% CI, 0.882\u20130.887) at 100 ms to 0.763 (0.758\u20130.769) at 800 ms for K = 10, showing retained local spatial topology even in the weakest cell (all networks positive; Holm-adjusted P = 7.63 \u00d7 10\u207b\u2076 across the four coefficient-free endpoints). e, Matched-minus-sequence-deranged centered-cosine similarity across the load-by-delay grid. f, Minimum matched-minus-sequence-deranged centered-cosine similarity across the 16 grid cells, 0.1337 (95% CI, 0.1298\u20130.1376), positive in all 20 networks, ruling out a common central template or smoothing artefact. n = 20 networks; two-sided 95% Student t CIs (bootstrap for the Holm family in d,f).')

# ---------------- S7 ----------------
h2('Supplementary Figure S7 | Conditional expression is robust to matching, window, definition, coverage and spatial-score controls')
figure(f'{S7_DIR}/supp_fig_s7.png', width_in=5.8)
legend_title('Supplementary Fig. S7 | The conditional effects in Fig. 7e,f are robust to exact area-and-energy matching, analysis window, parameter definition, coverage and spatial-score shuffling. ')
legend('a, Loss difference for the primary interaction on all trials and on the exact-match subset (2.61 [95% CI, 2.57\u20132.64] and 2.52 [2.48\u20132.56] percentage points), showing that the effect does not depend on trial inclusion. b, Percentage of trials satisfying the exact area-and-energy matching criterion (75.18% [95% CI, 74.46\u201375.90]). c, Interaction magnitude across early windows of 5, 10, 15 and 20 ms (11.78, 16.00, 18.62 and 16.15 percentage points), with the primary 10-ms endpoint reported in Fig. 7f. d, Interaction magnitude across cue-retention quantile q and overlap threshold. e, Complete-case coverage for the corresponding interaction cells. f, Three-endpoint estimation on complete paired rows for the primary interaction: the observed interaction was 16.00 percentage points (95% CI, 15.93\u201316.07), whereas the score-map shuffled control was 0.069 percentage points (\u22120.019 to 0.157). Their within-network difference was 15.93 percentage points (15.82\u201316.04; two-sided exact sign-flip P = 1.91 \u00d7 10\u207b\u2076), showing that the observed spatial-score arrangement contributed beyond shuffled score assignment. n = 20 networks; intervals are two-sided 95% Student t CIs unless stated otherwise.')

# ---------------- Supplementary Table S1 ----------------
p = h2('Supplementary Table S1 | Model, encoding and training parameters')
p.paragraph_format.page_break_before = True
para('Values are the actual configuration used for the 20-network ensemble and all main-text assays (sources: src/core/network.py, src/data/encoding.py, src/training/train_sdnn.py and the persisted ensemble run configuration).', size=9.5)
rows = [
 ('Architecture','Layers','3-layer feedforward spiking network'),
 ('Architecture','Layer 1','2 \u2192 30 feature maps; 5 \u00d7 5 kernel, stride 1, padding 2; 2 \u00d7 2 max-pooling'),
 ('Architecture','Layer 2','30 \u2192 150 feature maps; 3 \u00d7 3 kernel, stride 1, padding 2; 2 \u00d7 2 max-pooling'),
 ('Architecture','Layer 3','10 classes \u00d7 20 readout neurons; input spatial size 8'),
 ('Architecture','Competition','Top-k winner-take-all within layer: k = 5 (L1), 10 (L2), 1 (L3)'),
 ('Architecture','Inhibition','Lateral inhibition 20 mV (L1, L2) and 10 mV (L3); \u03c4inh = 10 ms'),
 ('Neuron','Passive properties','Vreset = \u221260 mV; VL = \u221270 mV; VE = 0 mV; Cm = 0.1 nF; gm = 10 nS; \u03c4e = 5 ms; refractory 20 ms; dt = 1 ms'),
 ('STSP','Utilization','U = 0.2 (baseline); u initialized to U; x initialized to 1'),
 ('STSP','Recovery','\u03c4D = 100 ms (depression); \u03c4F = 1,000 ms (facilitation)'),
 ('STSP','Training state','STSP disabled during training; enabled in post-training assays'),
 ('STSP','Gain compensation','Weights scaled by 1/U = 5 at load to preserve baseline drive'),
 ('Encoding','DoG preprocessing','ON/OFF kernels 7 \u00d7 7; \u03c31 = 1.0, \u03c32 = 2.0; threshold 0.05; per-channel max normalization'),
 ('Encoding','Latency','Rank-based latency within a 20-step window; stronger evidence \u2192 earlier spikes'),
 ('Encoding','Oscillatory schedule','Theta 5 Hz; gamma 50 Hz; active gamma indices 0, 3, 6'),
 ('Encoding','Input format','28 \u00d7 7 grayscale MNIST, resize and tensor conversion only'),
 ('Dataset','Split','MNIST 60,000 training / 10,000 test images; no validation split'),
 ('Training','Mini-batches','512 images per batch'),
 ('Training','Epochs','Layer 1: 2; Layer 2: 10; Layer 3: 100'),
 ('Training','Learning','L1/L2 learning rate 0.001; L3 learning rate 0.01; \u03c4+ = 20 ms; \u03c4elig = 20 ms'),
 ('Training','Weight bounds','wmin = 0, wmax = 1 nS; target trace level 0.5; reward/punishment maxima 1.0'),
 ('Training','Initial weights','L1/L2 0.6 nS; L3 0.8 nS'),
 ('Ensemble','Networks','20 independently trained networks, seeds 1000\u20131019'),
 ('Ensemble','Baseline accuracy','Mean test accuracy 91.158% (95% CI, 90.998\u201391.318%; range 90.76\u201391.90%)'),
 ('Readout','Decision gate','Earliest eligible Layer 3 spike at t mod 60 ms = 20 ms; no-response trials remain in analysis'),
]
table(['Group','Parameter','Value'], rows)

# ---------------- Supplementary Table S2 ----------------
p = h2('Supplementary Table S2 | Endpoint definitions and statistical analysis plan')
p.paragraph_format.page_break_before = True
para('One row per main-text inferential endpoint or endpoint family. Main-figure numbering follows the current seven-figure manuscript; deterministic Fig. 1 contains no network-level inferential endpoint. All endpoints use the independently trained network as the unit; trial-, anchor-, site-, pair- and stage-level values were aggregated within each network before cross-network inference. P values are the adjusted values reported in the main text and Source Data.', size=9.5)
rows2 = [
 ('Fig. 2a','Recall accuracy','Descriptive','\u2014','None','\u2014','Descriptive reference range 85\u201395%','\u2014'),
 ('Fig. 2c','Linear decoding accuracy','Descriptive','\u2014','None','\u2014','10% chance reference','\u2014'),
 ('Fig. 2d','Error-trial composition after u/x shuffling','Descriptive','\u2014','None','\u2014','Donor-item recomposition','\u2014'),
 ('Fig. 3b','Rescue: aligned minus mismatched history rate','S0-error anchors (baseline-incorrect B trials)','Separate opportunity sets','Two-sided one-sample t on network contrast','BH','0','0.000307'),
 ('Fig. 3b','Loss: aligned minus mismatched history rate','S0-correct anchors (baseline-correct B trials)','Separate opportunity sets','Two-sided one-sample t on network contrast','BH','0','0.000178'),
 ('Fig. 3c','Common input-driven component cosine similarity','Unmanipulated updates, margin above 0.5','200-ms B window','Two-sided one-sample t on margin','BH','0.5','3.64 \u00d7 10\u207b\u2075\u2075'),
 ('Fig. 3c','History-conditioned residual norm ratio','Margin above 0.05','200-ms B window','Two-sided one-sample t on margin','BH','0.05','1.53 \u00d7 10\u207b\u2074\u2075'),
 ('Fig. 3d','Residual magnitude at history-differential events','Size-matched random events (paired)','200-ms B window','Two-sided one-sample t on paired difference','BH','0','7.27 \u00d7 10\u207b\u2075\u2070'),
 ('Fig. 4a','Accuracy-drop contrast, dynamic STSP vs overlap-aligned reset','Non-overlap and size-matched random resets','First 50 ms','Two-sided one-sample t','BH','0','2.76 \u00d7 10\u207b\u00b9\u00b9'),
 ('Fig. 4a','Accuracy-drop contrast, non-overlap reset','Overlap-aligned reset reference','First 50 ms','Two-sided one-sample t','BH','0','2.76 \u00d7 10\u207b\u00b9\u00b9'),
 ('Fig. 4a','Accuracy-drop contrast, size-matched random reset','Overlap-aligned reset reference','First 50 ms','Two-sided one-sample t','BH','0','0.000780'),
 ('Fig. 4b','Pre-input support: overlap-dominant vs probe-only / balanced / random-matched','Probe-active pathway groups','Pre-input window','Two-sided one-sample t on contrasts','BH','0','2.05 \u00d7 10\u207b\u2079\u00b9 / 7.08 \u00d7 10\u207b\u2074\u2079 / 3.82 \u00d7 10\u207b\u00b9\u2075'),
 ('Fig. 4d','Dynamic-minus-attenuation and dynamic-minus-reset changes','Overlap-aligned attenuation/reset','First 50 ms','Two-sided one-sample t','BH','0','1.99 \u00d7 10\u207b\u00b3\u00b2 / 3.85 \u00d7 10\u207b\u00b3\u2076'),
 ('Fig. 4e','History-aligned updating difference-in-differences','Static-frozen control (no STSP mutation)','Dynamic vs static-frozen','Two-sided one-sample t','Unadjusted','0','1.24 \u00d7 10\u207b\u00b3\u2076'),
 ('Fig. 4f','Layer 2 donor-transfer index after Layer 1 u/x substitution','Identical B; receiver fast state equalized','Successor readout','Two-sided one-sample t','BH','0','2.44 \u00d7 10\u207b\u2074\u2079'),
 ('Fig. 5a','Early Layer 2 event-map donor-transfer index, K = 1 and K = 5','Transplanted successor; identical C','Early response to C','One-sided exact sign-flip over networks','Holm','0','3.81 \u00d7 10\u207b\u2076'),
 ('Fig. 5b','Post-C Layer 3 successor donor-transfer index, K = 1 and K = 5','Transplanted successor; identical C','Post-C successor','One-sided exact sign-flip over networks','Holm','0','3.81 \u00d7 10\u207b\u2076'),
 ('Fig. 5c','Observed-minus-passive joint-state displacement','Equal-time zero-input branch from same boundary','Stages 2\u201310','One-sided exact sign-flip over networks','Holm','0','1.91 \u00d7 10\u207b\u2075'),
 ('Fig. 5d','Rescue and loss shifts from K = 1 to K = 5','Relation-balanced opportunity sets','K = 1 vs K = 5','Two-sided exact sign-flip over networks','Holm','0','1.91 \u00d7 10\u207b\u2075'),
 ('Fig. 6b','Experienced-pair vs one-constituent-held shuffled pair cosine similarity','50 shuffled pairs per network','Terminal Layer 2 state','Two-sided one-sample t','BH','0','2.13 \u00d7 10\u207b\u00b3\u2078'),
 ('Fig. 7a','Cue-strength-integrated AUC gain, A and B vs no-memory and singleton','Cue-only no-memory reference; relevant singleton state','Cue-strength integration','Two-sided one-sample t','BH','0','\u22645.24 \u00d7 10\u207b\u2078'),
 ('Fig. 7b','Mean sequence-minus-singleton access gain','Slot-matched singleton states','Serial positions, K = 10, 400 ms','Two-sided one-sample t','BH','0','5.14 \u00d7 10\u207b\u00b9\u2079'),
 ('Fig. 7c','Matched vs same-label novel and unseen cues','Same-label novel; unseen-class cues','K = 7, 400 ms','Two-sided one-sample t','BH','0','1.60 \u00d7 10\u207b\u2076 / 1.03 \u00d7 10\u207b\u00b2\u2070'),
 ('Fig. 7d','Standardized sequence-length \u00d7 delay interaction','Within-network standardized product coefficient','4 \u00d7 4 load-delay grid','Two-sided one-sample t','Unadjusted','0','1.13 \u00d7 10\u207b\u00b9\u00b3'),
 ('Fig. 7e','High-overlap removal vs area- and energy-matched removal','Matched removal control','Recruitment loss','Two-sided one-sample t','BH','0','5.01 \u00d7 10\u207b\u00b3\u2070'),
 ('Fig. 7f','Overlap-gated STSP interaction (10 ms)','High vs low support \u00d7 overlap \u2265 0.05 vs 0','Primary 10-ms window','Two-sided one-sample t','BH','0','3.32 \u00d7 10\u207b\u2074\u2070'),
 ('Fig. S2b','Layer 1-only Layer 2 update and early Layer 3 donor transfer','Causal-identity gate (exact B, L1 only)','Transfer endpoints','One-sided exact sign-flip / confirmatory','Holm (core-8)','0','7.63 \u00d7 10\u207b\u2076'),
 ('Fig. S2d','Untouched-cohort transfer (n = 19)','Development seed 1000 excluded','Same endpoints','One-sided exact sign-flip','Holm','0','1.53 \u00d7 10\u207b\u2075'),
]
table(['Figure/panel','Endpoint','Reference / eligible set','Window / aggregation','Test','Family','Null','Adjusted P'], rows2, font=8)

para('', size=8)
para('Boundary notes. Supplementary figures protect the specific main-text claims listed in their legends and do not extend to other panels or conclusions. Panel-level engineering gates (identity audits, coverage, availability) are reported as descriptive values, not hypothesis tests. Post-hoc robustness analyses are two-sided by default and use exact sign-flip tests with Holm adjustment within each figure\u2019s declared family; no P value below the finite-enumeration bound (2\u207b\u00b2\u2070 = 9.54 \u00d7 10\u207b\u2077) is reported.', size=9.5)

cp = doc.core_properties
cp.title = 'Supplementary Information - History-Conditioned STSP Links Working-Memory Maintenance and Organization'
cp.language = 'en-US'

doc.save(OUT)
print('saved', OUT)
