from __future__ import annotations

import argparse
import os
from pathlib import Path

from docx import Document


def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Apply evidence-checked supplementary references and metadata corrections to the manuscript DOCX."
    )
    parser.add_argument("--input", type=Path, default=Path("docs/paper/v6.docx"))
    parser.add_argument("--output", type=Path, default=Path("docs/paper/v6.docx"))
    return parser.parse_args()


def _document_shape(document: Document) -> dict[str, int]:
    root = document.element.body
    return {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "inline_shapes": len(document.inline_shapes),
        "equations": len(root.xpath(".//m:oMath | .//m:oMathPara")),
        "drawings": len(root.xpath(".//w:drawing")),
    }


def _replace_in_run(document: Document, *, anchor: str, old: str, new: str) -> None:
    paragraphs = [paragraph for paragraph in document.paragraphs if anchor in paragraph.text]
    if len(paragraphs) != 1:
        raise ValueError(f"Expected one paragraph containing {anchor!r}, found {len(paragraphs)}")
    matching_runs = [run for run in paragraphs[0].runs if old in run.text]
    if len(matching_runs) != 1:
        raise ValueError(
            f"Expected one run containing {old!r} in paragraph anchored by {anchor!r}, found {len(matching_runs)}"
        )
    matching_runs[0].text = matching_runs[0].text.replace(old, new, 1)


def main() -> int:
    args = _parse_args()
    input_path = args.input.resolve()
    output_path = args.output.resolve()
    document = Document(input_path)
    before = _document_shape(document)

    replacements = [
        (
            "The reciprocal shift toward the donor item",
            "22.5%–28.4%; Fig. 2d).",
            "22.5%–28.4%; Fig. 2d; Supplementary Fig. S1).",
        ),
        (
            "Targeted attenuation or reset likewise",
            "; Fig. 4b–d),",
            "; Fig. 4b–d; Supplementary Fig. S3),",
        ),
        (
            "The Layer 2 successor shifted strongly",
            "; Fig. 4f).",
            "; Fig. 4f; Supplementary Fig. S2).",
        ),
        (
            "At both depths, the transplant shifted",
            "; Fig. 5a,b).",
            "; Fig. 5a,b; Supplementary Fig. S4).",
        ),
        (
            "Input-associated displacement exceeded passive evolution",
            "; Fig. 5c).",
            "; Fig. 5c; Supplementary Fig. S5).",
        ),
        (
            "Across sequence lengths K = 3, 5, 7, and 10",
            "; Fig. 6e,f).",
            "; Fig. 6e,f; Supplementary Fig. S6).",
        ),
        (
            "Removing the high-STSP-overlap contribution",
            "; Fig. 7f).",
            "; Fig. 7f; Supplementary Fig. S7).",
        ),
        (
            "mean post-training recall accuracy was 91.158%",
            "mean post-training recall accuracy was 91.158% (95% CI, 90.156–92.160%)",
            "mean post-training test accuracy was 91.158% (95% CI, 90.998–91.318%)",
        ),
        (
            "Source Data underlying Figs. 1–6",
            "Source Data underlying Figs. 1–6, ",
            "Source Data underlying Figs. 2–7 and Supplementary Figs. S1–S7, ",
        ),
        (
            "Source Data underlying Figs. 2–7",
            "provided with this manuscript. Analysis code",
            "provided with this manuscript. Fig. 1 is a deterministic illustration generated directly from the model equations and stated parameters. Analysis code",
        ),
    ]
    for anchor, old, new in replacements:
        _replace_in_run(document, anchor=anchor, old=old, new=new)

    after = _document_shape(document)
    if after != before:
        raise ValueError(f"DOCX structural counts changed unexpectedly: before={before}, after={after}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    temporary = output_path.with_name(f".{output_path.name}.tmp")
    document.save(temporary)
    os.replace(temporary, output_path)

    validated = Document(output_path)
    if _document_shape(validated) != before:
        raise ValueError("Saved DOCX failed structural-count validation")
    text = "\n".join(paragraph.text for paragraph in validated.paragraphs)
    required = [
        "Supplementary Fig. S1",
        "Supplementary Fig. S2",
        "Supplementary Fig. S3",
        "Supplementary Fig. S4",
        "Supplementary Fig. S5",
        "Supplementary Fig. S6",
        "Supplementary Fig. S7",
        "mean post-training test accuracy was 91.158% (95% CI, 90.998–91.318%)",
        "Source Data underlying Figs. 2–7 and Supplementary Figs. S1–S7",
    ]
    missing = [value for value in required if value not in text]
    if missing:
        raise ValueError(f"Saved DOCX is missing required text: {missing}")
    forbidden = ["90.156–92.160%", "Source Data underlying Figs. 1–6"]
    present = [value for value in forbidden if value in text]
    if present:
        raise ValueError(f"Saved DOCX retains forbidden text: {present}")

    print({"output": str(output_path), "structure": before, "replacements": len(replacements)})
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
