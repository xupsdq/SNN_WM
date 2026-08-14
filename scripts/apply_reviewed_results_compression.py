from __future__ import annotations

import argparse
import os
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


Chunk = tuple[str, str]


def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Apply the seven user-reviewed Results paragraph replacements to the manuscript DOCX."
    )
    parser.add_argument("--input", type=Path, default=Path("docs/paper/v6.docx"))
    parser.add_argument("--output", type=Path, default=Path("docs/paper/v6.docx"))
    return parser.parse_args()


def _document_shape(document: Document) -> dict[str, int]:
    root = document.element.body
    return {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "inline_shapes": len(document.inline_shapes),
        "equations": len(root.xpath(".//m:oMath | .//m:oMathPara")),
        "drawings": len(root.xpath(".//w:drawing")),
    }


def _clear_paragraph_content(paragraph: Paragraph) -> None:
    paragraph_xml = paragraph._p
    for child in list(paragraph_xml):
        if child.tag != paragraph_xml.pPr.tag:
            paragraph_xml.remove(child)


def _replace_paragraph(document: Document, *, old: str, chunks: list[Chunk]) -> None:
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text == old]
    if len(matches) != 1:
        raise ValueError(f"Expected one exact paragraph match, found {len(matches)}: {old[:120]!r}")
    paragraph = matches[0]
    _clear_paragraph_content(paragraph)
    for text, formatting in chunks:
        run = paragraph.add_run(text)
        if formatting == "italic":
            run.italic = True
        elif formatting == "superscript":
            run.font.superscript = True
        elif formatting != "plain":
            raise ValueError(f"Unknown formatting token: {formatting}")


def _text(chunks: list[Chunk]) -> str:
    return "".join(text for text, _ in chunks)


def main() -> int:
    args = _parse_args()
    input_path = args.input.resolve()
    output_path = args.output.resolve()
    document = Document(input_path)
    before = _document_shape(document)

    replacements: list[tuple[str, list[Chunk]]] = [
        (
            "To further distinguish functional inheritance from a merely decodable synaptic trace, we exchanged the retained joint u/x state between trials at readout while preserving the current input and fixed long-term weights. Among error trials, predictions of the original item fell from 23.6% (95% CI, 19.4%–27.8%) to 9.5% (7.4%–11.6%), whereas predictions of the donor item increased from 8.4% (5.9%–10.9%) to 25.5% (22.5%–28.4%; Fig. 2d; Supplementary Fig. S1). The reciprocal shift toward the donor item showed that later readout followed the substituted synaptic state, demonstrating that the retained STSP state was functionally inherited rather than merely decodable.",
            [
                ("To distinguish functional inheritance from a decodable trace, we exchanged retained joint ", "plain"),
                ("u/x", "italic"),
                (" states between trials at readout while holding the current input and fixed long-term weights constant. Among error trials, original-item predictions fell from 23.6% (95% CI, 19.4%–27.8%) to 9.5% (7.4%–11.6%), whereas donor-item predictions rose from 8.4% (5.9%–10.9%) to 25.5% (22.5%–28.4%; Fig. 2d). Dynamic–static differences weakened with delay, and paired flux and calibration excluded donor-opportunity bias (Supplementary Fig. S1). The donor-directed shift showed that later readout followed the substituted STSP state, establishing functional inheritance rather than mere decodability.", "plain"),
            ],
        ),
        (
            "To test whether inherited STSP supplied the history dependence, we selectively perturbed Layer 1 sites where retained effective support overlapped the incoming input. Resetting these overlap-aligned sites eliminated the accuracy cost relative to the static-frozen control: compared with this reset, accuracy dropped 5.60% more under both intact dynamic STSP and a non-overlap reset (95% CI, 4.76%–6.44%; BH-adjusted P = 2.76 × 10−11) and 1.34% more after a size-matched random reset (95% CI, 0.64%–2.04%; BH-adjusted P = 0.000780; Fig. 4a). Targeted attenuation or reset likewise reduced the history-associated advancement and recruitment of early spikes (n = 20 networks; mean first-50-ms dynamic-minus-attenuation change, 7.45 percentage points [95% CI, 7.36–7.53], BH-adjusted P = 1.99 × 10−32; dynamic-minus-reset change, 37.27 percentage points [37.01–37.53], BH-adjusted P = 3.85 × 10−36; Fig. 4b–d; Supplementary Fig. S3), while dynamic STSP produced stronger history-aligned downstream updating than the static-frozen control (difference-in-differences, 8.52%; 95% CI, 8.46%–8.58%; P = 1.24 × 10−36; Fig. 4e). These results localized the history-dependent effect to inherited STSP engaged by the incoming input.",
            [
                ("To localize history dependence, we perturbed Layer 1 sites where retained effective support overlapped the incoming input. Resetting these sites abolished the accuracy cost relative to the static-frozen control: intact dynamics and a non-overlap reset each reduced accuracy by an additional 5.60% (95% CI, 4.76%–6.44%; BH-adjusted ", "plain"),
                ("P", "italic"),
                (" = 2.76 × 10", "plain"),
                ("−11", "superscript"),
                ("), and a size-matched random reset by 1.34% (0.64%–2.04%; BH-adjusted ", "plain"),
                ("P", "italic"),
                (" = 0.000780; Fig. 4a). Attenuation and reset reduced early-spike advancement and recruitment by 7.45 percentage points (95% CI, 7.36–7.53; BH-adjusted ", "plain"),
                ("P", "italic"),
                (" = 1.99 × 10", "plain"),
                ("−32", "superscript"),
                (") and 37.27 percentage points (37.01–37.53; BH-adjusted ", "plain"),
                ("P", "italic"),
                (" = 3.85 × 10", "plain"),
                ("−36", "superscript"),
                ("), respectively (Fig. 4b–d). These effects were robust to window, winner-rank and distance definitions, and the interventions chiefly delayed rather than eliminated original winners (Supplementary Fig. S3). Dynamic STSP also strengthened history-aligned downstream updating relative to the static-frozen control (difference-in-differences, 8.52%; 95% CI, 8.46%–8.58%; ", "plain"),
                ("P", "italic"),
                (" = 1.24 × 10", "plain"),
                ("−36", "superscript"),
                ("; Fig. 4e). Thus, history dependence localized to inherited STSP recruited along input-overlapping pathways.", "plain"),
            ],
        ),
        (
            "Localization identified the source of history dependence but did not establish whether inherited STSP directed successor formation. We therefore exchanged only the inherited Layer 1 u/x state between donor and receiver histories while holding the current input, fixed long-term weights, other retained STSP states and fast state constant. The Layer 2 successor shifted strongly toward the donor history (donor-transfer index, 0.8086; 95% CI, 0.8075–0.8098; BH-adjusted P = 2.44 × 10−49; Fig. 4f; Supplementary Fig. S2). Thus, the inherited STSP state was sufficient to redirect successor formation toward the transplanted history.",
            [
                ("Localization did not establish whether inherited STSP directed successor formation. We therefore exchanged only the Layer 1 ", "plain"),
                ("u/x", "italic"),
                (" state between histories while holding the current input, fixed long-term weights, other retained STSP states and fast state constant. The Layer 2 successor shifted strongly toward the donor history (donor-transfer index, 0.8086; 95% CI, 0.8075–0.8098; BH-adjusted ", "plain"),
                ("P", "italic"),
                (" = 2.44 × 10", "plain"),
                ("−49", "superscript"),
                ("; Fig. 4f). Early Layer 3 transfer and metric and cohort sensitivities confirmed that this effect extended beyond the transplanted layer (Supplementary Fig. S2). Thus, inherited Layer 1 STSP was sufficient to redirect successor formation under an identical current input.", "plain"),
            ],
        ),
        (
            "To test whether a successor could itself condition the next input, we transplanted the post-B Layer 2 u/x state from a donor history into a matched receiver while preserving the receiver’s other retained states and fast state, and then presented the identical input C. We tested both shallow and deeper history depths (K = 1 and K = 5). At both depths, the transplant shifted the early Layer 2 response to C toward the donor history (donor-transfer index, 0.337 [95% bootstrap CI, 0.336–0.339] and 0.344 [0.342–0.345], respectively) and likewise shifted the post-C Layer 3 successor (0.283 [0.281–0.284] and 0.291 [0.290–0.292]; all Holm-adjusted P = 3.81 × 10−6; Fig. 5a,b; Supplementary Fig. S4). Thus, the successor formed by one transition was sufficient to shape both processing of the next input and the state formed from it.",
            [
                ("To test whether a successor conditioned the next input, we transplanted the post-B Layer 2 ", "plain"),
                ("u/x", "italic"),
                (" state between matched histories, held other retained and fast states fixed, and presented the identical input C. At ", "plain"),
                ("K", "italic"),
                (" = 1 and ", "plain"),
                ("K", "italic"),
                (" = 5, the transplant shifted both the early Layer 2 response toward the donor history (donor-transfer index, 0.337 [95% bootstrap CI, 0.336–0.339] and 0.344 [0.342–0.345]) and the post-C Layer 3 successor (0.283 [0.281–0.284] and 0.291 [0.290–0.292]; all Holm-adjusted ", "plain"),
                ("P", "italic"),
                (" = 3.81 × 10", "plain"),
                ("−6", "superscript"),
                ("; Fig. 5a,b). Exact identity checks and exclusion of the development seed preserved these effects (Supplementary Fig. S4). Thus, a successor was sufficient to condition both the next input and the state formed from it.", "plain"),
            ],
        ),
        (
            "Selective transplantation established that a successor was sufficient to condition the next transition, but did not show that the same updating recurred during unmanipulated sequences. We therefore compared the joint u/x displacement produced by each successive input with equal-time passive evolution from the same preceding boundary. Input-associated displacement exceeded passive evolution at every tested boundary, with a mean observed-minus-passive displacement of 0.532 (95% bootstrap CI, 0.528–0.536; Holm-adjusted P = 1.91 × 10−5; Fig. 5c; Supplementary Fig. S5). Over the same history range, from K = 1 to K = 5, rescue decreased by 23.1% (95% bootstrap CI, 19.2%–26.5%) and loss increased by 33.7% (32.6%–34.8%; both Holm-adjusted P = 1.91 × 10−5; Fig. 5d). Thus, input-associated state updating recurred across successive inputs beyond passive STSP evolution, while interference increased with history depth.",
            [
                ("Transplantation established successor sufficiency but not recurrence during unmanipulated sequences. We therefore compared joint ", "plain"),
                ("u/x", "italic"),
                (" displacement after each input with equal-time passive evolution from the same preceding boundary. Input-associated displacement exceeded passive evolution at every tested boundary (mean difference, 0.532; 95% bootstrap CI, 0.528–0.536; Holm-adjusted ", "plain"),
                ("P", "italic"),
                (" = 1.91 × 10", "plain"),
                ("−5", "superscript"),
                ("; Fig. 5c). The excess remained positive at each network’s weakest stage and involved both ", "plain"),
                ("u", "italic"),
                (" and ", "plain"),
                ("x", "italic"),
                (" (Supplementary Fig. S5). From ", "plain"),
                ("K", "italic"),
                (" = 1 to ", "plain"),
                ("K", "italic"),
                (" = 5, rescue decreased by 23.1% (95% bootstrap CI, 19.2%–26.5%) and loss increased by 33.7% (32.6%–34.8%; both Holm-adjusted ", "plain"),
                ("P", "italic"),
                (" = 1.91 × 10", "plain"),
                ("−5", "superscript"),
                ("; Fig. 5d). Thus, input-driven STSP updating recurred as interference accumulated.", "plain"),
            ],
        ),
        (
            "We next asked whether this organization persisted as history accumulated. Across sequence lengths K = 3, 5, 7, and 10, effective component number increased while the contribution of the latest item declined (mean effective component number at K = 3, 5, 7 and 10, respectively: 2.994 [95% CI, 2.993–2.995], 4.957 [4.952–4.962], 6.743 [6.723–6.763] and 8.048 [7.913–8.182]; mean latest-item weight at the same loads: 0.331 [0.330–0.332], 0.195 [0.194–0.197], 0.133 [0.132–0.134] and 0.0726 [0.0714–0.0738]; Fig. 6c,d), indicating increasingly distributed component expression rather than latest-item domination. Across the same load–delay conditions, history-matched STSP morphology remained more similar to the terminal state than sequence-deranged composites (across-condition mean effective STSP-support area, 0.444 [K = 3, 800 ms; 95% CI, 0.437–0.451] to 0.566 [K = 10, 100 ms; 0.560–0.571]; mean matched-minus-deranged centered-cosine difference, 0.134 [K = 10, 100 ms; 0.130–0.138] to 0.387 [K = 3, 800 ms; 0.375–0.399]; Fig. 6e,f; Supplementary Fig. S6). Thus, longer histories remained distributed across multiple components while preserving history-specific spatial organization.",
            [
                ("We next asked whether this organization persisted as history accumulated. From ", "plain"),
                ("K", "italic"),
                (" = 3 to ", "plain"),
                ("K", "italic"),
                (" = 10, effective component number increased from 2.994 (95% CI, 2.993–2.995) to 8.048 (7.913–8.182), while latest-item weight fell from 0.331 (0.330–0.332) to 0.0726 (0.0714–0.0738; Fig. 6c,d). Across load and delay, matched STSP morphology remained more similar to the terminal state than sequence-deranged composites (centered-cosine difference, 0.134–0.387; Fig. 6e,f). Normalized-coefficient and coefficient-free analyses preserved this distributed, locally organized and history-specific pattern (Supplementary Fig. S6). Thus, accumulated histories remained distributed without losing their spatial organization.", "plain"),
            ],
        ),
        (
            "We next asked whether functional expression of retained STSP required the incoming input to engage the same synaptic pathways. Removing the high-STSP-overlap contribution produced 2.52% more recruitment loss than area- and energy-matched removal (95% CI, 2.48%–2.56%; BH-adjusted P = 5.01 × 10−30; Fig. 7e). Consistently, retained STSP altered early firing only along overlapping pathways, yielding an overlap-gated interaction of 16.0% (95% CI, 15.9%–16.1%; BH-adjusted P = 3.32 × 10−40; Fig. 7f; Supplementary Fig. S7).",
            [
                ("We next asked whether retained STSP was expressed through pathways engaged by the incoming input. Removing high-overlap contributions caused 2.52% more recruitment loss than area- and energy-matched removal (95% CI, 2.48%–2.56%; BH-adjusted ", "plain"),
                ("P", "italic"),
                (" = 5.01 × 10", "plain"),
                ("−30", "superscript"),
                ("; Fig. 7e). Retained STSP altered early firing along overlapping pathways, producing a 16.0% interaction (95% CI, 15.9%–16.1%; BH-adjusted ", "plain"),
                ("P", "italic"),
                (" = 3.32 × 10", "plain"),
                ("−40", "superscript"),
                ("; Fig. 7f). The effect persisted under exact matching and across tested definitions, and exceeded complete-paired spatial-score shuffling (Supplementary Fig. S7). Thus, incoming input accessed retained STSP through pathway overlap.", "plain"),
            ],
        ),
    ]

    for old, chunks in replacements:
        _replace_paragraph(document, old=old, chunks=chunks)

    after = _document_shape(document)
    if after != before:
        raise ValueError(f"DOCX structural counts changed unexpectedly: before={before}, after={after}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    temporary = output_path.with_name(f".{output_path.name}.tmp")
    document.save(temporary)
    os.replace(temporary, output_path)

    validated = Document(output_path)
    if _document_shape(validated) != before:
        raise ValueError("Saved DOCX failed structural-count validation")
    full_text = "\n".join(paragraph.text for paragraph in validated.paragraphs)
    expected = [_text(chunks) for _, chunks in replacements]
    missing = [text[:100] for text in expected if text not in full_text]
    if missing:
        raise ValueError(f"Saved DOCX is missing reviewed paragraph(s): {missing}")
    forbidden = [old for old, _ in replacements]
    retained = [text[:100] for text in forbidden if text in full_text]
    if retained:
        raise ValueError(f"Saved DOCX retains replaced paragraph(s): {retained}")

    print({"output": str(output_path), "structure": before, "paragraph_replacements": len(replacements)})
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
